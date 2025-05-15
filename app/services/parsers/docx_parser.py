# app/services/parsers/docx_parser.py

from docx import Document

def parse_docx(file_path: str) -> str:
    """
    Extracts text and table contents from a DOCX file.
    """
    doc = Document(file_path)
    extracted_text = ""

    # Extract paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            extracted_text += para.text.strip() + "\n"

    # Extract tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells)
            extracted_text += row_text + "\n"

    return extracted_text.strip()
